"""
============================================================================
Document Text Extraction Script (with OCR & Image Support)
============================================================================
This script extracts text from various document formats:
- PDF files (.pdf) - including scanned PDFs via OCR
- Word documents (.docx, .doc)
- Text files (.txt)
- Markdown files (.md)
- Images (.png, .jpg, .jpeg) - via OCR

NEW: Supports scanned PDFs and can extract/describe images!

USAGE:
    python scripts/extract_text.py
    
    With OCR enabled (for scanned documents):
    python scripts/extract_text.py --ocr
    
    Extract and describe images:
    python scripts/extract_text.py --ocr --extract_images

REQUIREMENTS:
    pip install PyPDF2 pdfplumber python-docx chardet pytesseract pdf2image Pillow tqdm rich
    
    For OCR, you also need Tesseract installed:
    - Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki
    - Linux: sudo apt install tesseract-ocr
    - Mac: brew install tesseract
============================================================================
"""

import os
import sys
import re
import argparse
from pathlib import Path
from typing import Optional, List, Tuple
from dataclasses import dataclass
import logging

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from tqdm import tqdm
from rich.console import Console
from rich.logging import RichHandler

# Setup rich console for beautiful output
console = Console()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(message)s",
    handlers=[RichHandler(console=console, rich_tracebacks=True)]
)
logger = logging.getLogger(__name__)


# ============================================================================
# OCR FUNCTIONS
# ============================================================================
def check_tesseract_installed() -> bool:
    """Check if Tesseract OCR is installed."""
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def ocr_image(image, language: str = 'eng') -> str:
    """
    Perform OCR on an image.
    
    Args:
        image: PIL Image object
        language: OCR language (default: English)
        
    Returns:
        Extracted text
    """
    try:
        import pytesseract
        text = pytesseract.image_to_string(image, lang=language)
        return text.strip()
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def pdf_page_to_image(pdf_path: Path, page_num: int):
    """
    Convert a PDF page to an image for OCR.
    
    Args:
        pdf_path: Path to PDF file
        page_num: Page number (0-indexed)
        
    Returns:
        PIL Image object or None
    """
    try:
        from pdf2image import convert_from_path
        
        # Convert just the one page
        images = convert_from_path(
            pdf_path,
            first_page=page_num + 1,
            last_page=page_num + 1,
            dpi=200  # Good balance of quality and speed
        )
        
        if images:
            return images[0]
        return None
    except Exception as e:
        logger.warning(f"Could not convert PDF page to image: {e}")
        return None


def is_scanned_page(page_text: str, threshold: int = 50) -> bool:
    """
    Check if a PDF page is likely scanned (has very little extractable text).
    
    Args:
        page_text: Text extracted from the page
        threshold: Minimum character count to consider as "has text"
        
    Returns:
        True if page appears to be scanned
    """
    # Remove whitespace and count actual characters
    clean_text = ''.join(page_text.split())
    return len(clean_text) < threshold


# ============================================================================
# IMAGE EXTRACTION FROM PDF
# ============================================================================
def extract_images_from_pdf(pdf_path: Path, output_dir: Path) -> List[Path]:
    """
    Extract images from a PDF file.
    
    Args:
        pdf_path: Path to PDF file
        output_dir: Directory to save images
        
    Returns:
        List of paths to extracted images
    """
    import fitz  # PyMuPDF - we'll add this as optional
    
    extracted_images = []
    
    try:
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Save image
                    image_name = f"{pdf_path.stem}_p{page_num + 1}_img{img_idx + 1}.{image_ext}"
                    image_path = output_dir / image_name
                    
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    
                    extracted_images.append(image_path)
                    
                except Exception as e:
                    logger.debug(f"Could not extract image: {e}")
        
        doc.close()
        
    except ImportError:
        logger.warning("PyMuPDF not installed. Cannot extract embedded images.")
        logger.info("Install with: pip install PyMuPDF")
    except Exception as e:
        logger.warning(f"Could not extract images from {pdf_path.name}: {e}")
    
    return extracted_images


# ============================================================================
# PDF EXTRACTION (with OCR support)
# ============================================================================
def extract_text_from_pdf(
    pdf_path: Path,
    use_ocr: bool = False,
    ocr_language: str = 'eng'
) -> Tuple[str, dict]:
    """
    Extract text from a PDF file, with optional OCR for scanned pages.
    
    Args:
        pdf_path: Path to the PDF file
        use_ocr: Whether to use OCR for scanned pages
        ocr_language: Language for OCR
        
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    text_parts = []
    metadata = {
        "source": str(pdf_path),
        "pages": 0,
        "extractor": None,
        "ocr_pages": 0
    }
    
    # Try pdfplumber first (handles complex layouts better)
    try:
        import pdfplumber
        
        with pdfplumber.open(pdf_path) as pdf:
            metadata["pages"] = len(pdf.pages)
            metadata["extractor"] = "pdfplumber"
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                
                # Check if this is a scanned page
                if use_ocr and is_scanned_page(page_text):
                    logger.debug(f"Page {page_num} appears scanned, using OCR...")
                    image = pdf_page_to_image(pdf_path, page_num - 1)
                    if image:
                        ocr_text = ocr_image(image, ocr_language)
                        if ocr_text:
                            page_text = ocr_text
                            metadata["ocr_pages"] += 1
                
                if page_text:
                    text_parts.append(f"\n[Page {page_num}]\n{page_text}")
                    
        if text_parts:
            return "\n".join(text_parts), metadata
            
    except Exception as e:
        logger.warning(f"pdfplumber failed for {pdf_path.name}: {e}")
    
    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(pdf_path)
        metadata["pages"] = len(reader.pages)
        metadata["extractor"] = "PyPDF2"
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text() or ""
            
            # Check if this is a scanned page
            if use_ocr and is_scanned_page(page_text):
                image = pdf_page_to_image(pdf_path, page_num - 1)
                if image:
                    ocr_text = ocr_image(image, ocr_language)
                    if ocr_text:
                        page_text = ocr_text
                        metadata["ocr_pages"] += 1
            
            if page_text:
                text_parts.append(f"\n[Page {page_num}]\n{page_text}")
                
    except Exception as e:
        logger.error(f"PyPDF2 also failed for {pdf_path.name}: {e}")
        raise
    
    return "\n".join(text_parts), metadata


# ============================================================================
# IMAGE TEXT EXTRACTION (OCR)
# ============================================================================
def extract_text_from_image(image_path: Path, language: str = 'eng') -> Tuple[str, dict]:
    """
    Extract text from an image file using OCR.
    
    Args:
        image_path: Path to the image file
        language: OCR language
        
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    from PIL import Image
    
    metadata = {"source": str(image_path), "type": "image", "ocr": True}
    
    try:
        image = Image.open(image_path)
        text = ocr_image(image, language)
        return text, metadata
    except Exception as e:
        logger.error(f"Failed to OCR image {image_path.name}: {e}")
        return "", metadata


# ============================================================================
# WORD DOCUMENT EXTRACTION
# ============================================================================
def extract_text_from_docx(docx_path: Path) -> Tuple[str, dict]:
    """
    Extract text from a Word document (.docx).
    
    Args:
        docx_path: Path to the Word document
        
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    from docx import Document
    
    doc = Document(docx_path)
    metadata = {"source": str(docx_path), "paragraphs": len(doc.paragraphs)}
    
    text_parts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    return "\n\n".join(text_parts), metadata


# ============================================================================
# TEXT FILE EXTRACTION
# ============================================================================
def extract_text_from_txt(txt_path: Path) -> Tuple[str, dict]:
    """
    Extract text from a plain text file with automatic encoding detection.
    
    Args:
        txt_path: Path to the text file
        
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    import chardet
    
    # Detect encoding
    with open(txt_path, 'rb') as f:
        raw_data = f.read()
        detected = chardet.detect(raw_data)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
    
    metadata = {"source": str(txt_path), "encoding": encoding}
    
    # Read with detected encoding
    try:
        text = raw_data.decode(encoding)
    except UnicodeDecodeError:
        # Fallback to utf-8 with error handling
        text = raw_data.decode('utf-8', errors='ignore')
        metadata["encoding_fallback"] = True
    
    return text, metadata


# ============================================================================
# TEXT CLEANING
# ============================================================================
def clean_text(text: str, aggressive: bool = False) -> str:
    """
    Clean extracted text by removing noise and normalizing formatting.
    
    Args:
        text: Raw extracted text
        aggressive: If True, apply more aggressive cleaning
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Normalize unicode characters
    import unicodedata
    text = unicodedata.normalize('NFKC', text)
    
    # Replace common problematic characters
    replacements = {
        '\x00': '',           # Null bytes
        '\r\n': '\n',         # Windows line endings
        '\r': '\n',           # Old Mac line endings
        '\t': ' ',            # Tabs to spaces
        '\xa0': ' ',          # Non-breaking spaces
        '\u200b': '',         # Zero-width spaces
        '\ufeff': '',         # BOM
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Remove multiple consecutive spaces
    text = re.sub(r' +', ' ', text)
    
    # Remove multiple consecutive newlines (keep max 2)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    if aggressive:
        # Remove lines that are just numbers (page numbers)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Remove lines that are too short (likely noise)
        lines = [line for line in text.split('\n') if len(line) > 3 or line == '']
        text = '\n'.join(lines)
        
        # Remove repeated punctuation
        text = re.sub(r'([.!?])\1+', r'\1', text)
    
    # Final strip
    text = text.strip()
    
    return text


# ============================================================================
# MAIN EXTRACTION FUNCTION
# ============================================================================
@dataclass
class ExtractionResult:
    """Container for extraction results."""
    filename: str
    success: bool
    text: str
    metadata: dict
    error: Optional[str] = None
    

def extract_document(
    file_path: Path,
    clean: bool = True,
    use_ocr: bool = False,
    ocr_language: str = 'eng'
) -> ExtractionResult:
    """
    Extract text from a document file.
    
    Args:
        file_path: Path to the document
        clean: Whether to clean the extracted text
        use_ocr: Whether to use OCR for scanned content
        ocr_language: Language for OCR
        
    Returns:
        ExtractionResult with the extracted text and metadata
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    try:
        # Choose extractor based on file type
        if suffix == '.pdf':
            text, metadata = extract_text_from_pdf(file_path, use_ocr, ocr_language)
        elif suffix in ['.docx', '.doc']:
            text, metadata = extract_text_from_docx(file_path)
        elif suffix in ['.txt', '.md']:
            text, metadata = extract_text_from_txt(file_path)
        elif suffix in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif']:
            if use_ocr:
                text, metadata = extract_text_from_image(file_path, ocr_language)
            else:
                return ExtractionResult(
                    filename=file_path.name,
                    success=False,
                    text="",
                    metadata={"source": str(file_path)},
                    error="Image files require --ocr flag"
                )
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        # Clean text if requested
        if clean:
            text = clean_text(text)
        
        metadata["char_count"] = len(text)
        metadata["word_count"] = len(text.split())
        
        return ExtractionResult(
            filename=file_path.name,
            success=True,
            text=text,
            metadata=metadata
        )
        
    except Exception as e:
        logger.error(f"Failed to extract {file_path.name}: {e}")
        return ExtractionResult(
            filename=file_path.name,
            success=False,
            text="",
            metadata={"source": str(file_path)},
            error=str(e)
        )


# ============================================================================
# BATCH PROCESSING
# ============================================================================
def process_documents(
    input_dir: Path,
    output_dir: Path,
    extensions: List[str] = None,
    clean: bool = True,
    use_ocr: bool = False,
    ocr_language: str = 'eng'
) -> List[ExtractionResult]:
    """
    Process all documents in a directory.
    
    Args:
        input_dir: Directory containing documents
        output_dir: Directory to save extracted text
        extensions: List of file extensions to process
        clean: Whether to clean extracted text
        use_ocr: Whether to use OCR
        ocr_language: Language for OCR
        
    Returns:
        List of ExtractionResult objects
    """
    if extensions is None:
        extensions = ['.pdf', '.docx', '.doc', '.txt', '.md']
        if use_ocr:
            extensions.extend(['.png', '.jpg', '.jpeg', '.bmp', '.tiff'])
    
    input_dir = Path(input_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Find all documents
    documents = []
    for ext in extensions:
        documents.extend(input_dir.glob(f"**/*{ext}"))
    
    if not documents:
        logger.warning(f"No documents found in {input_dir}")
        return []
    
    # =====================================================
    # DEBUG: Show all documents found
    # =====================================================
    console.print("\n[bold cyan]═══════════════════════════════════════════════════════════════[/bold cyan]")
    console.print("[bold cyan]                    📄 DOCUMENT EXTRACTION DEBUG                    [/bold cyan]")
    console.print("[bold cyan]═══════════════════════════════════════════════════════════════[/bold cyan]\n")
    
    console.print(f"[bold green]✓ Found {len(documents)} documents to extract:[/bold green]\n")
    
    # Group by file type
    by_type = {}
    for doc in documents:
        ext = doc.suffix.lower()
        by_type.setdefault(ext, []).append(doc)
    
    for ext, files in sorted(by_type.items()):
        console.print(f"  [yellow]{ext}[/yellow]: {len(files)} file(s)")
        for f in files:
            file_size = f.stat().st_size
            size_str = f"{file_size:,} bytes" if file_size < 1024 else f"{file_size/1024:.1f} KB"
            console.print(f"    [dim]•[/dim] {f.name} [dim]({size_str})[/dim]")
    
    console.print()
    
    if use_ocr:
        if check_tesseract_installed():
            console.print("[green]✓ Tesseract OCR is available[/green]")
        else:
            console.print("[yellow]⚠ Tesseract OCR not found. OCR will be skipped.[/yellow]")
            logger.info("Install Tesseract:")
            logger.info("  Windows: https://github.com/UB-Mannheim/tesseract/wiki")
            logger.info("  Linux: sudo apt install tesseract-ocr")
            logger.info("  Mac: brew install tesseract")
            use_ocr = False
    
    console.print()
    
    results = []
    successful = 0
    failed = 0
    ocr_count = 0
    
    # Process each document with detailed output
    console.print("[bold yellow]Extracting text from documents...[/bold yellow]\n")
    
    for doc_path in tqdm(documents, desc="Extracting text"):
        result = extract_document(doc_path, clean=clean, use_ocr=use_ocr, ocr_language=ocr_language)
        results.append(result)
        
        if result.success:
            successful += 1
            
            # Track OCR usage
            used_ocr = result.metadata.get("ocr_pages", 0) > 0 or result.metadata.get("ocr")
            if used_ocr:
                ocr_count += 1
            
            # Save extracted text
            output_filename = doc_path.stem + ".txt"
            output_path = output_dir / output_filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result.text)
            
            # Save metadata
            metadata_path = output_dir / (doc_path.stem + "_metadata.txt")
            with open(metadata_path, 'w', encoding='utf-8') as f:
                for key, value in result.metadata.items():
                    f.write(f"{key}: {value}\n")
            
            # Debug output for each file
            word_count = result.metadata.get('word_count', 0)
            ocr_tag = " [OCR]" if used_ocr else ""
            console.print(f"  [green]✓[/green] {doc_path.name}{ocr_tag}: {word_count:,} words extracted")
        else:
            failed += 1
            console.print(f"  [red]✗[/red] {doc_path.name}: {result.error}")
    
    # =====================================================
    # Summary table
    # =====================================================
    console.print("\n[bold cyan]═══════════════════════════════════════════════════════════════[/bold cyan]")
    console.print("[bold cyan]                    📊 EXTRACTION SUMMARY                         [/bold cyan]")
    console.print("[bold cyan]═══════════════════════════════════════════════════════════════[/bold cyan]\n")
    
    console.print(f"[bold green]✓ Successfully extracted: {successful} documents[/bold green]")
    if ocr_count > 0:
        console.print(f"[bold blue]  └─ {ocr_count} documents used OCR[/bold blue]")
    if failed > 0:
        console.print(f"[bold red]✗ Failed: {failed} documents[/bold red]")
    
    console.print(f"\n[dim]Output saved to: {output_dir}[/dim]")
    
    return results


# ============================================================================
# COMMAND LINE INTERFACE
# ============================================================================
def main():
    """Main entry point for the extraction script."""
    parser = argparse.ArgumentParser(
        description="Extract text from documents (PDF, DOCX, TXT, images)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    # Basic extraction (text-based PDFs only)
    python scripts/extract_text.py
    
    # With OCR for scanned documents
    python scripts/extract_text.py --ocr
    
    # Process image files too
    python scripts/extract_text.py --ocr --extensions .pdf .png .jpg
    
    # Specify OCR language (e.g., German)
    python scripts/extract_text.py --ocr --ocr_language deu
        """
    )
    
    # Import config for defaults
    from config import DOCUMENTS_DIR, DATA_DIR
    
    parser.add_argument(
        "--input_dir",
        type=Path,
        default=DOCUMENTS_DIR,
        help="Directory containing documents to process"
    )
    
    parser.add_argument(
        "--output_dir",
        type=Path,
        default=DATA_DIR / "extracted",
        help="Directory to save extracted text"
    )
    
    parser.add_argument(
        "--extensions",
        nargs="+",
        default=None,
        help="File extensions to process (default: .pdf .docx .doc .txt .md)"
    )
    
    parser.add_argument(
        "--no-clean",
        action="store_true",
        help="Skip text cleaning"
    )
    
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="Enable OCR for scanned documents and images"
    )
    
    parser.add_argument(
        "--ocr_language",
        type=str,
        default="eng",
        help="OCR language (default: eng). Use 'eng+fra' for multiple."
    )
    
    args = parser.parse_args()
    
    console.print("[bold blue]Document Text Extraction[/bold blue]")
    console.print(f"Input directory: {args.input_dir}")
    console.print(f"Output directory: {args.output_dir}")
    if args.ocr:
        console.print(f"[green]OCR enabled[/green] (language: {args.ocr_language})")
    console.print()
    
    # Check input directory exists
    if not args.input_dir.exists():
        console.print(f"[red]Error: Input directory does not exist: {args.input_dir}[/red]")
        console.print("[yellow]Please add your PDF and document files to this directory.[/yellow]")
        args.input_dir.mkdir(parents=True, exist_ok=True)
        console.print(f"[green]Created directory: {args.input_dir}[/green]")
        return
    
    # Process documents
    results = process_documents(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        extensions=args.extensions,
        clean=not args.no_clean,
        use_ocr=args.ocr,
        ocr_language=args.ocr_language
    )
    
    # Print summary
    if results:
        total_words = sum(r.metadata.get("word_count", 0) for r in results if r.success)
        console.print(f"\n[bold]Total words extracted: {total_words:,}[/bold]")


if __name__ == "__main__":
    main()
